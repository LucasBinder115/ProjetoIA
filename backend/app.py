# === Imports ===
from flask import Flask, render_template, send_file, request, jsonify
import os
import requests
import nltk
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Image, PageBreak
from reportlab.lib.units import inch
from docx import Document
from ebooklib import epub
from functools import lru_cache
from celery import Celery
from redis import Redis
import logging
from celery.signals import after_setup_logger

# === Celery Setup ===
def make_celery(flask_app):
    celery_app = Celery(
        flask_app.import_name,
        broker='redis://localhost:6379/0',
        backend='redis://localhost:6379/0'
    )
    celery_app.conf.update(flask_app.config)
    class ContextTask(celery_app.Task):
        def __call__(self, *args, **kwargs):
            with flask_app.app_context():
                return self.run(*args, **kwargs)
    celery_app.Task = ContextTask
    return celery_app

# === Initial Setup ===
nltk.download('punkt', quiet=True)
app = Flask(__name__, template_folder='../templates', static_folder='../static')
app.config['UPLOAD_FOLDER'] = 'static/imagens'
celery = make_celery(app)

# Move this line to ensure celery is fully initialized before task definition
celery.conf.update(task_track_started=True)  # Optional: track task start

redis = Redis(host='localhost', port=6379)

# === Celery Configuration ===
@after_setup_logger.connect
def setup_loggers(logger, *args, **kwargs):
    logger.addHandler(logging.FileHandler('celery.log'))

# === Helper Functions ===
@lru_cache(maxsize=100)
def gerar_html(questoes):
    pass

def buscar_questoes(tema):
    url = f"https://opentdb.com/api.php?amount=10&category=18&type=multiple"
    response = requests.get(url)
    if response.status_code == 200:
        results = response.json()['results']
        return [{"question": q["question"]} for q in results]
    return [{"question": "Erro ao buscar questões"}]

# === File Creation Functions ===
def criar_pdf(caminho_imagem=None, questoes=None):
    html_content = """
    <!DOCTYPE html>
    <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial; margin: 2cm; }
                h1 { color: #2c3e50; }
                .questao { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <h1>Ebook de Questões para Estudo</h1>
    """
    if questoes:
        for i, questao in enumerate(questoes, start=1):
            html_content += f'<div class="questao"><strong>{i}.</strong> {questao["question"]}</div>'
    html_content += "</body></html>"

    files = {
        'files': ('index.html', html_content)
    }
    data = {
        'marginTop': '1',
        'marginBottom': '1',
        'marginLeft': '1',
        'marginRight': '1'
    }

    try:
        response = requests.post(
            'http://localhost:3000/forms/chromium/convert/html',
            files=files,
            data=data,
            stream=True,
            timeout=10
        )
        response.raise_for_status()
        output_path = 'prova.pdf'
        with open(output_path, 'wb') as f:
            f.write(response.content)
        logging.info(f"PDF gerado com sucesso em {os.path.abspath(output_path)}")
        return output_path
    except requests.exceptions.RequestException as e:
        logging.error(f"Erro na geração do PDF com Gotenberg: {e}")
        raise

def criar_docx(caminho_imagem=None, questoes=None):
    doc = Document()
    doc.add_heading('Ebook de Questões para Estudo', 0)
    if questoes:
        for i, questao in enumerate(questoes, start=1):
            doc.add_paragraph(f"{i}. {questao['question']}")
    doc.save('prova.docx')
    return 'prova.docx'

def criar_epub(caminho_imagem=None, questoes=None):
    livro = epub.EpubBook()
    livro.set_title('Ebook de Questões para Estudo')
    livro.set_language('pt-BR')
    if questoes:
        capitulo = epub.EpubHtml(title='Questões', file_name='questoes.xhtml', lang='pt-BR')
        capitulo.content = "<h1>Questões</h1><ul>"
        for questao in questoes:
            capitulo.content += f"<li>{questao['question']}</li>"
        capitulo.content += "</ul>"
        livro.add_item(capitulo)
        livro.toc = (epub.Link('questoes.xhtml', 'Questões', 'questoes'),)
    livro.add_item(epub.EpubNav())
    livro.add_item(epub.EpubNcx())
    epub.write_epub('prova.epub', livro)
    return 'prova.epub'

# === Celery Tasks ===
@celery.task(name='app.gerar_pdf_async')  # Explicitly name the task
def gerar_pdf_async(prompt, formato, caminho_imagem=None):
    questoes = buscar_questoes(prompt if prompt else "default")

    if formato == "pdf":
        caminho_arquivo = criar_pdf(caminho_imagem, questoes)
    elif formato == "docx":
        caminho_arquivo = criar_docx(caminho_imagem, questoes)
    elif formato == "epub":
        caminho_arquivo = criar_epub(caminho_imagem, questoes)
    else:
        logging.error(f"Formato inválido recebido: {formato}")
        raise ValueError(f"Formato inválido: {formato}. Use 'pdf', 'docx', ou 'epub'.")

    cache_key = f"pdf:{prompt}:{formato}"
    redis.setex(cache_key, 3600, caminho_arquivo)
    return caminho_arquivo

# === Flask Routes ===
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/escolher_formato')
def escolher_formato():
    return render_template('escolher_formato.html')

@app.route('/gerar_ebook', methods=['POST'])
def gerar_ebook():
    formato = request.form.get('formato')
    prompt = request.form.get('prompt')
    imagem = request.files.get('imagem')

    valid_formats = {"pdf", "docx", "epub"}
    if not formato or formato not in valid_formats:
        return jsonify({"error": "Formato inválido. Use 'pdf', 'docx', ou 'epub'."}), 400

    caminho_imagem = None
    if imagem:
        caminho_imagem = os.path.join(app.config['UPLOAD_FOLDER'], imagem.filename)
        imagem.save(caminho_imagem)

    task = gerar_pdf_async.delay(prompt, formato, caminho_imagem)
    return jsonify({"task_id": task.id}), 202

@app.route('/status/<task_id>')
def check_status(task_id):
    task = gerar_pdf_async.AsyncResult(task_id)
    if task.successful():
        caminho_arquivo = task.result
        return jsonify({
            "status": task.status,
            "download_url": f"/download/{os.path.basename(caminho_arquivo)}"
        })
    return jsonify({"status": task.status, "download_url": None})

@app.route('/download/<filename>')
def download_file(filename):
    caminho_arquivo = os.path.join(os.getcwd(), filename)
    if os.path.exists(caminho_arquivo):
        return send_file(caminho_arquivo, as_attachment=True)
    return jsonify({"error": "Arquivo não encontrado"}), 404

@app.route('/health')
def health_check():
    try:
        response = requests.get('http://localhost:3000/health')
        response.raise_for_status()
        return "OK", 200
    except requests.exceptions.RequestException:
        return "Gotenberg offline", 500

@app.route('/test_ia')
def test_ia():
    try:
        questoes = buscar_questoes("Química Orgânica")
        if not isinstance(questoes, list):
            raise ValueError("Questões não são uma lista")
        questoes_serializaveis = [{"question": q["question"]} for q in questoes]
        return jsonify(questoes_serializaveis)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# === Test Functions ===
def test_gotenberg():
    questoes = [{'question': 'Teste 1'}, {'question': 'Teste 2'}]
    criar_pdf(questoes=questoes)
    print("PDF gerado com sucesso!")

# === Main Execution ===
if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    app.run(debug=True)