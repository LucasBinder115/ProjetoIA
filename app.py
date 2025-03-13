# === Imports ===
from flask import Flask, render_template, send_file, request
import os
import requests
import nltk
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Image, PageBreak
from reportlab.lib.units import inch
from docx import Document
from ebooklib import epub
from transformers import pipeline

# === Initial Setup ===
nltk.download('punkt')
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/imagens'  # Folder for uploaded images

# === Helper Functions ===
def buscar_questoes(tema):
    url = f"https://opentdb.com/api.php?amount=10&category=18&type=multiple"
    response = requests.get(url)
    return response.json()['results'] if response.status_code == 200 else []

def generate_question(prompt):
    gerador = pipeline('text-generation', model='EleutherAI/gpt-neo-1.3B')
    return gerador(prompt, max_length=50)[0]['generated_text']

# === PDF Creation Functions ===
def criar_pdf(caminho_imagem=None, questoes=None):
    pdf = SimpleDocTemplate("prova.pdf", pagesize=letter)
    styles = getSampleStyleSheet()
    conteudo = []

    # Title
    conteudo.append(Paragraph("Ebook de Questões para Estudo", styles['Title']))
    conteudo.append(Spacer(1, 12))

    # Add uploaded image if provided
    if caminho_imagem:
        try:
            imagem = Image(caminho_imagem, width=4*inch, height=3*inch)
            conteudo.append(imagem)
            conteudo.append(Spacer(1, 12))
        except Exception as e:
            print(f"Erro ao carregar a imagem: {e}")
            conteudo.append(Paragraph("<i>Imagem não disponível</i>", styles['Italic']))
            conteudo.append(Spacer(1, 12))

    # Add questions if provided
    if questoes:
        for i, questao in enumerate(questoes, start=1):
            questao_formatada = Paragraph(f"{i}. {questao['question']}", styles['BodyText'])
            conteudo.append(questao_formatada)
            conteudo.append(Spacer(1, 12))
    else:
        # Default themed questions
        questoes_por_tema = {
            "Matemática": {
                "imagem": "static/imagens/matematica.jpg",
                "questoes": [
                    "1. Qual é a fórmula para calcular a área de um círculo?",
                    "2. Explique o Teorema de Pitágoras."
                ]
            },
            "História": {
                "imagem": "static/imagens/historia.jpg",
                "questoes": [
                    "1. Quais foram as causas da Revolução Francesa?",
                    "2. Descreva o período da Revolução Industrial."
                ]
            },
            "Biologia": {
                "imagem": "static/imagens/biologia.jpg",
                "questoes": [
                    "1. O que é fotossíntese e qual a sua importância?",
                    "2. Explique a estrutura do DNA."
                ]
            }
        }

        for tema, dados in questoes_por_tema.items():
            conteudo.append(Paragraph(f"<b>Tema: {tema}</b>", styles['Heading2']))
            conteudo.append(Spacer(1, 12))
            
            try:
                imagem = Image(dados["imagem"], width=4*inch, height=3*inch)
                conteudo.append(imagem)
                conteudo.append(Spacer(1, 12))
            except Exception as e:
                print(f"Erro ao carregar a imagem: {e}")
                conteudo.append(Paragraph("<i>Imagem não disponível</i>", styles['Italic']))
                conteudo.append(Spacer(1, 12))

            for questao in dados["questoes"]:
                conteudo.append(Paragraph(questao, styles['BodyText']))
                conteudo.append(Spacer(1, 12))
            
            conteudo.append(PageBreak())

    pdf.build(conteudo)

# === DOCX Creation Functions ===
def criar_docx(caminho_imagem=None, questoes=None):
    doc = Document()
    doc.add_heading('Ebook de Questões para Estudo', 0)

    if questoes:
        for i, questao in enumerate(questoes, start=1):
            doc.add_paragraph(f"{i}. {questao['question']}")
    else:
        questoes_por_tema = {
            "Matemática": [
                "1. Qual é a fórmula para calcular a área de um círculo?",
                "2. Explique o Teorema de Pitágoras."
            ],
            "História": [
                "1. Quais foram as causas da Revolução Francesa?",
                "2. Descreva o período da Revolução Industrial."
            ],
            "Biologia": [
                "1. O que é fotossíntese e qual a sua importância?",
                "2. Explique a estrutura do DNA."
            ]
        }

        for tema, questoes_list in questoes_por_tema.items():
            doc.add_heading(f'Tema: {tema}', level=1)
            for questao in questoes_list:
                doc.add_paragraph(questao)
            doc.add_page_break()

    doc.save('prova.docx')

# === EPUB Creation Functions ===
def criar_epub(caminho_imagem=None, questoes=None):
    livro = epub.EpubBook()
    livro.set_title('Ebook de Questões para Estudo')
    livro.set_language('pt-BR')

    if questoes:
        capitulo = epub.EpubHtml(title='Questões', file_name='questoes.xhtml', lang='pt-BR')
        capitulo.content = "<h1>Questões</h1><ul>"
        for questao in questoes:
            capitulo.content += f"<li>{questao['question']}</li>"
        capitulo.content += "</ul>"
        livro.add_item(capitulo)
        livro.toc = (epub.Link('questoes.xhtml', 'Questões', 'questoes'),)
    else:
        questoes_por_tema = {
            "Matemática": [
                "1. Qual é a fórmula para calcular a área de um círculo?",
                "2. Explique o Teorema de Pitágoras."
            ],
            "História": [
                "1. Quais foram as causas da Revolução Francesa?",
                "2. Descreva o período da Revolução Industrial."
            ],
            "Biologia": [
                "1. O que é fotossíntese e qual a sua importância?",
                "2. Explique a estrutura do DNA."
            ]
        }

        for tema, questoes_list in questoes_por_tema.items():
            capitulo = epub.EpubHtml(title=tema, file_name=f'{tema}.xhtml', lang='pt-BR')
            capitulo.content = f"<h1>{tema}</h1><ul>"
            for questao in questoes_list:
                capitulo.content += f"<li>{questao}</li>"
            capitulo.content += "</ul>"
            livro.add_item(capitulo)

        livro.toc = tuple(questoes_por_tema.keys())

    livro.add_item(epub.EpubNav())
    livro.add_item(epub.EpubNcx())
    epub.write_epub('prova.epub', livro)

# === Flask Routes ===
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/escolher_formato')
def escolher_formato():
    return render_template('escolher_formato.html')

@app.route('/gerar_ebook', methods=['POST'])
def gerar_ebook():
    formato = request.form.get('formato')
    tema = request.form.get('tema')
    imagem = request.files.get('imagem')
    
    caminho_imagem = None
    if imagem:
        caminho_imagem = os.path.join(app.config['UPLOAD_FOLDER'], imagem.filename)
        imagem.save(caminho_imagem)

    questoes = buscar_questoes(tema) if tema else None

    if formato == "pdf":
        criar_pdf(caminho_imagem, questoes)
        return send_file('prova.pdf', as_attachment=True)
    elif formato == "docx":
        criar_docx(caminho_imagem, questoes)
        return send_file('prova.docx', as_attachment=True)
    elif formato == "epub":
        criar_epub(caminho_imagem, questoes)
        return send_file('prova.epub', as_attachment=True)
    else:
        return "Formato inválido."

# === Main Execution ===
if __name__ == '__main__':
    app.run(debug=True)